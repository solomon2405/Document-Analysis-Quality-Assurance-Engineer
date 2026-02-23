from __future__ import annotations

import asyncio
import hashlib
import logging
import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

import pdfplumber
from docx import Document
from fastapi import UploadFile
from openpyxl import load_workbook

from app.config import settings
from app.services.ocr_service import OCRResult, run_ocr

logger = logging.getLogger(__name__)
TOKEN_PATTERN = re.compile(r"\S+")
DATE_PATTERN = re.compile(r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b")
CURRENCY_PATTERN = re.compile(r"\$[\d,]+(?:\.\d{1,2})?")
PERCENT_PATTERN = re.compile(r"\b\d+(?:\.\d+)?%\b")
SECTION_REF_PATTERN = re.compile(r"\bsection\s+\d+(?:\.\d+)?\b", re.IGNORECASE)
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg"}


@dataclass(slots=True)
class TextUnit:
    text: str
    file: str
    page: int
    paragraph: int
    line: int
    heading_level: int = 0
    section_key: str = ""


@dataclass(slots=True)
class OCRImageDoc:
    file: str
    words: list[TextUnit]
    blocks: list[dict[str, object]]
    avg_confidence: float
    fingerprint: str


@dataclass(slots=True)
class UnifiedDocument:
    side: str
    units: list[TextUnit]
    full_text: str
    headings: list[TextUnit]
    section_map: dict[str, list[TextUnit]]
    tables: list[TextUnit]
    bullets: list[TextUnit]
    image_docs: list[OCRImageDoc]
    entities: dict[str, set[str]]
    numeric_values: set[str]
    section_references: set[str]
    source_files: list[str]


def _tokenize(line: str, file_name: str, page: int, paragraph: int, line_no: int) -> list[TextUnit]:
    return [
        TextUnit(
            text=m.group(0),
            file=file_name,
            page=page,
            paragraph=paragraph,
            line=line_no,
        )
        for m in TOKEN_PATTERN.finditer(line)
    ]


def _is_heading(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    if len(stripped) <= 90 and stripped == stripped.upper():
        return True
    return bool(re.match(r"^\d+(\.\d+)*\s+\w+", stripped))


def _extract_entities_and_numbers(text: str, entities: dict[str, set[str]], numerics: set[str], refs: set[str]) -> None:
    for token in re.findall(r"\b\d+(?:[.,]\d+)?\b", text):
        numerics.add(token)
    for match in DATE_PATTERN.findall(text):
        entities["DATE"].add(match)
    for match in CURRENCY_PATTERN.findall(text):
        entities["MONEY"].add(match)
        numerics.add(match)
    for match in PERCENT_PATTERN.findall(text):
        entities["PERCENT"].add(match)
        numerics.add(match)
    for match in SECTION_REF_PATTERN.finditer(text):
        refs.add(match.group(0).strip().lower())

    for word in re.findall(r"\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\b", text):
        if len(word) > 2:
            entities["ENTITY"].add(word)


def _extract_docx(file_name: str, file_bytes: bytes) -> tuple[list[TextUnit], list[TextUnit], list[TextUnit]]:
    doc = Document(BytesIO(file_bytes))
    units: list[TextUnit] = []
    headings: list[TextUnit] = []
    bullets: list[TextUnit] = []
    for para_idx, para in enumerate(doc.paragraphs, start=1):
        line = (para.text or "").strip()
        if not line:
            continue
        row_tokens = _tokenize(line, file_name, 1, para_idx, para_idx)
        if _is_heading(line):
            for t in row_tokens:
                t.heading_level = 1
                t.section_key = line[:80]
            headings.extend(row_tokens)
        if line.startswith(("-", "*")):
            bullets.extend(row_tokens)
        units.extend(row_tokens)
    return units, headings, bullets


def _extract_pdf(file_name: str, file_bytes: bytes) -> tuple[list[TextUnit], list[TextUnit], list[TextUnit]]:
    units: list[TextUnit] = []
    headings: list[TextUnit] = []
    bullets: list[TextUnit] = []
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page_no, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            for line_no, line in enumerate(text.splitlines(), start=1):
                clean = line.strip()
                if not clean:
                    continue
                row_tokens = _tokenize(clean, file_name, page_no, line_no, line_no)
                if _is_heading(clean):
                    for t in row_tokens:
                        t.heading_level = 1
                        t.section_key = clean[:80]
                    headings.extend(row_tokens)
                if clean.startswith(("-", "*")):
                    bullets.extend(row_tokens)
                units.extend(row_tokens)
    return units, headings, bullets


def _extract_txt(file_name: str, file_bytes: bytes) -> tuple[list[TextUnit], list[TextUnit], list[TextUnit]]:
    units: list[TextUnit] = []
    headings: list[TextUnit] = []
    bullets: list[TextUnit] = []
    for line_no, line in enumerate(file_bytes.decode("utf-8", errors="ignore").splitlines(), start=1):
        clean = line.strip()
        if not clean:
            continue
        row_tokens = _tokenize(clean, file_name, 1, line_no, line_no)
        if _is_heading(clean):
            for t in row_tokens:
                t.heading_level = 1
                t.section_key = clean[:80]
            headings.extend(row_tokens)
        if clean.startswith(("-", "*")):
            bullets.extend(row_tokens)
        units.extend(row_tokens)
    return units, headings, bullets


def _extract_xlsx(file_name: str, file_bytes: bytes) -> tuple[list[TextUnit], list[TextUnit], list[TextUnit], list[TextUnit]]:
    wb = load_workbook(filename=BytesIO(file_bytes), read_only=True, data_only=True)
    units: list[TextUnit] = []
    tables: list[TextUnit] = []
    try:
        para_no = 0
        for sheet in wb.worksheets:
            for row_no, row in enumerate(sheet.iter_rows(values_only=True), start=1):
                values = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
                if not values:
                    continue
                para_no += 1
                line = " | ".join(values)
                row_tokens = _tokenize(line, file_name, 1, para_no, row_no)
                units.extend(row_tokens)
                tables.extend(row_tokens)
    finally:
        wb.close()
    return units, [], [], tables


def _extract_image(file_name: str, file_bytes: bytes) -> tuple[list[TextUnit], OCRImageDoc]:
    ocr: OCRResult = run_ocr(file_bytes)
    words = [
        TextUnit(
            text=w.text,
            file=file_name,
            page=1,
            paragraph=w.line_no,
            line=w.line_no,
        )
        for w in ocr.words
    ]
    digest = hashlib.sha256(file_bytes).hexdigest()[:16]
    blocks = [
        {"text": b.text, "bbox": b.bbox, "line": b.line_no, "block": b.block_no}
        for b in ocr.blocks
    ]
    return words, OCRImageDoc(
        file=file_name,
        words=words,
        blocks=blocks,
        avg_confidence=ocr.avg_confidence,
        fingerprint=digest,
    )


def _extract_single(file_name: str, file_bytes: bytes) -> dict[str, object]:
    suffix = Path(file_name).suffix.lower()
    if suffix not in settings.allowed_suffixes:
        raise ValueError(f"Unsupported file type: {suffix}")

    if suffix == ".docx":
        units, headings, bullets = _extract_docx(file_name, file_bytes)
        return {"units": units, "headings": headings, "bullets": bullets, "tables": [], "images": []}
    if suffix == ".pdf":
        units, headings, bullets = _extract_pdf(file_name, file_bytes)
        return {"units": units, "headings": headings, "bullets": bullets, "tables": [], "images": []}
    if suffix == ".txt":
        units, headings, bullets = _extract_txt(file_name, file_bytes)
        return {"units": units, "headings": headings, "bullets": bullets, "tables": [], "images": []}
    if suffix == ".xlsx":
        units, headings, bullets, tables = _extract_xlsx(file_name, file_bytes)
        return {"units": units, "headings": headings, "bullets": bullets, "tables": tables, "images": []}

    units, image_doc = _extract_image(file_name, file_bytes)
    return {"units": units, "headings": [], "bullets": [], "tables": [], "images": [image_doc]}


async def ingest_files(files: list[UploadFile], side: str) -> UnifiedDocument:
    ordered = sorted(files, key=lambda f: (f.filename or "").lower())

    async def process(file: UploadFile) -> dict[str, object]:
        raw = await file.read()
        if len(raw) > settings.max_file_size_mb * 1024 * 1024:
            raise ValueError(f"File {file.filename} exceeds size limit")
        extracted = await asyncio.to_thread(_extract_single, file.filename or "unknown", raw)
        extracted["file"] = file.filename or "unknown"
        return extracted

    results = await asyncio.gather(*(process(f) for f in ordered))

    units: list[TextUnit] = []
    headings: list[TextUnit] = []
    tables: list[TextUnit] = []
    bullets: list[TextUnit] = []
    images: list[OCRImageDoc] = []
    section_map: dict[str, list[TextUnit]] = {}
    entities = {"DATE": set(), "MONEY": set(), "PERCENT": set(), "ENTITY": set()}
    numerics: set[str] = set()
    section_refs: set[str] = set()
    source_files: list[str] = []

    for result in results:
        source_files.append(str(result["file"]))
        file_units = list(result["units"])
        units.extend(file_units)
        headings.extend(list(result["headings"]))
        bullets.extend(list(result["bullets"]))
        tables.extend(list(result["tables"]))
        images.extend(list(result["images"]))
        paragraph_groups: dict[tuple[int, int], list[TextUnit]] = {}
        for unit in file_units:
            paragraph_groups.setdefault((unit.page, unit.paragraph), []).append(unit)
        for _, grouped_units in paragraph_groups.items():
            paragraph_text = " ".join(u.text for u in grouped_units).strip().lower()
            key = grouped_units[0].section_key or paragraph_text[:120] or f"paragraph-{grouped_units[0].paragraph}"
            section_map.setdefault(key, []).extend(grouped_units)

        _extract_entities_and_numbers(" ".join(u.text for u in file_units), entities, numerics, section_refs)

    full_text = " ".join(u.text for u in units)
    return UnifiedDocument(
        side=side,
        units=units,
        full_text=full_text,
        headings=headings,
        section_map=section_map,
        tables=tables,
        bullets=bullets,
        image_docs=images,
        entities=entities,
        numeric_values=numerics,
        section_references=section_refs,
        source_files=source_files,
    )
